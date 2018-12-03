#x,y=input 2 letter words
#adjust width and height to store 100 barcodes in a single Page 
#Rows = 2500 for 100 pages
#pageNumber = str(x)+str(y)+"00"+str(i)+".png" for i in range(input())

import barcode
import img2pdf
from barcode.writer import ImageWriter
import time
from fpdf import FPDF
from PIL import Image
import docx
from docx.api import Document
import os
from time import time

start_time=time()
folderAddress="/home/zorro/blockchain/"

def makeFolder(folderAddress,directoryName):
    try:
        os.mkdir(folderAddress+directoryName)
    except FileExistsError:
        pass
        
Code_128="code128"

letter_words_1=input()
letter_words_2=input()

def fasterSolution(folderAddress,pic):
    document = Document()
    COLUMNS=4
    table=document.add_table(rows=1000,columns=COLUMNS)
    table_cells=table._cells
    for i in range(ROWS):
        row_cells=table_cells[i*COLUMNS:(i+1)*COLUMNS]
        for cell in row_cells.cells:
            paragraph=cell.paragraphs[0]
            run=paragraph.add_run()
            run.add_picture(folderAddress+pic+".png",width=350000*1.42,height=700000*0.49)
        
def putIntoDocumentFiles(folderAddress,pic):
    document = Document()
    table=document.add_table(rows=25,cols=4)#ROWS=25 For 100 barcodes 
    for row in table.rows:
        for cell in row.cells:
            paragraph=cell.paragraphs[0]
            run=paragraph.add_run()
            run.add_picture(folderAddress+pic+".png",width=350000*0.71,height=350000*0.49)#(width,height)=>dimensions(singleTableRow,singleTableColumn) for singlePage

    for i in range(2):
        document.save(folderAddress+"singleFolder"+"/"+pic+"_"+str(i+1)+".docx")
    

pageNumber=int(input())
array,array2=[],[]
for i in range(1,1+pageNumber):
    barcode_Name=letter_words_1+letter_words_2+"00"+str(i)
    ean=barcode.get(Code_128,barcode_Name,writer=ImageWriter())
    filename=ean.save(folderAddress+barcode_Name)#saves to [MG001.png,MG002.png...MG00(pageNumber).png]
    array.append(folderAddress+barcode_Name+".png")
    array2.append(barcode_Name)
    
#time.sleep(2)

makeFolder(folderAddress,"singleFolder1")
for a,b in zip(array,array2):
    #time.sleep(0.2)
    putIntoDocumentFiles(folderAddress,b)
    
end_time=time()
time_taken = end_time - start_time # time_taken is in seconds
hours, rest = divmod(time_taken,3600)
minutes, seconds = divmod(rest, 60)
