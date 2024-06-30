import barcode
import img2pdf
from barcode.writer import ImageWriter
import time
from fpdf import FPDF
from PIL import Image
import docx
from docx.api import Document
import os

folderAddress="/home/zorro/blockchain/"

def makeFolder(folderAddress,directoryName):
    try:
        os.mkdir(folderAddress+directoryName)
    except FileExistsError:
        pass
        
Code_128="code128"

letter_words_1=input()
letter_words_2=input()

def putIntoDocumentFiles(folderAddress,pic):
    document=docx.Document()
    table=document.add_table(rows=10,cols=4)
    
    for row in table.rows:
        for cell in row.cells:
            paragraph=cell.paragraphs[0]
            run=paragraph.add_run()
            run.add_picture(folderAddress+pic+".png",width=350000*1.42,height=700000*0.49)

    for i in range(5):#0+4+1+5+1+10+1
        document.save(folderAddress+"/"+pic+"/"+pic+"_"+str(i+1)+".docx")
    

def fd(image,savedName):
    width,height=image.size
    image=image.resize(int(width/2),int(height/2))
    image.save(savedName)
    
pageNumber=int(input())
array,array2=[],[]
for i in range(1,1+pageNumber):
    barcode_Name=letter_words_1+letter_words_2+"00"+str(i)
    ean=barcode.get(Code_128,barcode_Name,writer=ImageWriter())
    filename=ean.save(folderAddress+barcode_Name)#saves to [MG001.png,MG002.png...MG00(pageNumber).png]
    array.append(folderAddress+barcode_Name+".png")
    array2.append(barcode_Name)
    

time.sleep(2)

for a,b in zip(array,array2):
    #img=Image.open(a)#opens the image
    #fd(img,a)#resized the image and saved the image
    makeFolder(folderAddress,b)#a='/home/zorro/blockchain/'+''=>a[23:-4]+'.png'
    time.sleep(0.2)#take a break
    #putIntoDocumentFiles(folderAddress,a[23:-4])
    putIntoDocumentFiles(folderAddress,b)


