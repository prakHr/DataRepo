#######################################################################################
# x,y= 2 strings ######################################################################
# pageNumber = integer input which gives x+y+"00"+str(page)=> (1,length(pageNumber)) ##
# create images using png format and then resized images ##############################
# then saves into folders using corresponding names of x+y+"00"+str(page) #############
#######################################################################################
import barcode
import img2pdf
from barcode.writer import ImageWriter
import time
from fpdf import FPDF
from PIL import Image
import docx
from docx.api import Document
import os

folderAddress="/home/zorro/blockchain/"

def makeFolder(folderAddress,directoryName):
    try:
        os.mkdir(folderAddress+directoryName)
    except FileExistsError:
        pass

Code_128="code128"

letter_words_1=input()
letter_words_2=input()

def putIntoDocumentFiles(folderAddress,pic):
    document=docx.Document()
    table=document.add_table(rows=10,cols=4)

    for row in table.rows:
        for cell in row.cells:
            paragraph=cell.paragraphs[0]
            run=paragraph.add_run()
            run.add_picture(pic+".png",width=350000*1.42,height=700000*0.49)

    for i in range(100):
        document.save(folderAddress+"docxDocuments"+pic[22:]+"_"+str(i+1)+".docx")


def fd(image,savedName):
    width,height=image.size
    image=image.resize(int(width/2),int(height/2))
    image.save(savedName)

pageNumber=int(input())
array=[]
for i in range(1,1+pageNumber):
    barcode_Name=letter_words_1+letter_words_2+"00"+str(i)
    ean=barcode.get(Code_128,barcode_Name,writer=ImageWriter())
    filename=ean.save(barcode_Name)
    array.append(folderAddress+barcode_Name+".png")


time.sleep(3)

for a in array:
    makeFolder(folderAddress,a[23:-4])
    putIntoDocumentFiles(folderAddress,a[:-4])
